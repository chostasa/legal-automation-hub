import streamlit as st
st.set_page_config(page_title="Legal Automation Hub", layout="wide")

import pandas as pd
import os
import zipfile
import io
import tempfile
from docx import Document
st.markdown("""
<style>
.stButton > button {
    background-color: #B08B48;
    color: white;
    border: none;
    padding: 0.5rem 1rem;
    font-weight: 600;
}
.stTextInput > div > input {
    border: 1px solid #0A1D3B;
}
.stTextArea > div > textarea {
    border: 1px solid #0A1D3B;
}
</style>
""", unsafe_allow_html=True)


# === Simple login ===
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    password = st.text_input("Enter Password", type="password")
    if password == st.secrets["password"]:
        st.session_state.authenticated = True
        st.rerun()
    else:
        st.stop()

# === Branding: Logo inside navy header bar ===
import base64

# Load the logo as base64 to embed directly (Cloud-safe)
def load_logo_base64(file_path):
    with open(file_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode()

logo_base64 = load_logo_base64("sggh_logo.png")

st.markdown(f"""
<div style="background-color: #0A1D3B; padding: 2rem 0; text-align: center;">
    <img src="data:image/png;base64,{logo_base64}" width="360" style="margin-bottom: 1rem;" />
    <h1 style="color: white; font-size: 2.2rem; margin: 0;">Stinar Gould Grieco & Hensley</h1>
</div>
""", unsafe_allow_html=True)


# === Sidebar Navigation - Only visible after login ===
with st.sidebar:
    st.markdown("### ⚖️ Legal Automation Hub")
    tool = st.radio("Choose Tool", [  # ✅ inside the sidebar
        "📖 Instructions & Support",
        "📄 Batch Doc Generator",
        "📬 FOIA Requests",
        "📂 Demands",
        "🚧 Complaint (In Progress)",
        "🚧 Subpoenas (In Progress)",
    ])


# === Routing ===
if tool == "📂 Demands":
    st.header("📁 Generate Demand Letters")
    with st.form("demand_form"):
        client_name = st.text_input("Client Name")
        defendant = st.text_input("Defendant")
        incident_date = st.date_input("Incident Date")
        location = st.text_input("Location")
        summary = st.text_area("Summary of Incident")
        damages = st.text_area("Damages")

        submitted = st.form_submit_button("Generate Demand Letter")

    if submitted:
        df = pd.DataFrame([{
            "Client Name": client_name,
            "Defendant": defendant,
            "IncidentDate": incident_date.strftime("%B %d, %Y"),
            "Location": location,
            "Summary": summary,
            "Damages": damages
        }])

        try:
            output_paths = run_demand(df)  # <-- make sure this function is defined/imported
            st.success("✅ Letter generated!")
            for path in output_paths:
                filename = os.path.basename(path)
                with open(path, "rb") as f:
                    st.download_button(label=f"Download {filename}", data=f, file_name=filename)
        except Exception as e:
            st.error(f"❌ Error: {e}")

elif tool == "📬 FOIA Requests":
    st.header("📨 Generate FOIA Letters")
    with st.form("foia_form"):
        client_id = st.text_input("Client ID")
        defendant_name = st.text_input("Defendant Name")
        abbreviation = st.text_input("Defendant Abbreviation (for file name)")
        address_line1 = st.text_input("Defendant Address Line 1")
        address_line2 = st.text_input("Defendant Address Line 2 (City, State, Zip)")
        date_of_incident = st.date_input("Date of Incident")
        location = st.text_input("Location of Incident")
        case_synopsis = st.text_area("Case Synopsis")
        potential_requests = st.text_area("Potential Requests")
        explicit_instructions = st.text_area("Explicit Instructions (optional)")
        case_type = st.text_input("Case Type")
        facility = st.text_input("Facility or System")
        defendant_role = st.text_input("Defendant Role")

        submitted = st.form_submit_button("Generate FOIA Letter")

    if submitted:
        df = pd.DataFrame([{
            "Client ID": client_id,
            "Defendant Name": defendant_name,
            "Defendant Abbreviation": abbreviation,
            "Defendant Line 1 (address)": address_line1,
            "Defendant Line 2 (City,state, zip)": address_line2,
            "DOI": date_of_incident,
            "location of incident": location,
            "Case Synopsis": case_synopsis,
            "Potential Requests": potential_requests,
            "Explicit instructions": explicit_instructions,
            "Case Type": case_type,
            "Facility or System": facility,
            "Defendant Role": defendant_role
        }])

        try:
            output_paths = run_foia(df)  # <-- make sure this function is defined/imported
            st.success("✅ FOIA letter generated!")
            for path in output_paths:
                filename = os.path.basename(path)
                with open(path, "rb") as f:
                    st.download_button(label=f"Download {filename}", data=f, file_name=filename)
        except Exception as e:
            st.error(f"❌ Error: {e}")

elif tool == "📄 Batch Doc Generator":
    st.header("📄 Batch Document Generator")

    TEMPLATE_FOLDER = os.path.join("templates", "batch_docs")
    os.makedirs(TEMPLATE_FOLDER, exist_ok=True)

    # Upload new template
    st.subheader("📁 Upload a New Template")
    uploaded_template = st.file_uploader("Upload a .docx Template", type="docx")
    if uploaded_template:
        save_path = os.path.join(TEMPLATE_FOLDER, uploaded_template.name)
        with open(save_path, "wb") as f:
            f.write(uploaded_template.read())
        st.success(f"✅ Saved '{uploaded_template.name}' to your template library.")
        st.rerun()

    # Select saved template
    st.subheader("📂 Select a Saved Template")
    excluded_templates = {"foia_template.docx", "demand_template.docx"}
    available_templates = [
        f for f in os.listdir(TEMPLATE_FOLDER)
        if f.endswith(".docx") and f not in excluded_templates
    ]

    if not available_templates:
        st.warning("⚠️ No saved templates found. Upload one above.")
        st.stop()

    template_choice = st.selectbox("Choose Template", available_templates)
    template_path = os.path.join(TEMPLATE_FOLDER, template_choice)

    # Upload Excel and filename format
    excel_file = st.file_uploader("Upload Excel Data (.xlsx)", type="xlsx")
    output_name_format = st.text_input("Enter filename format (e.g., HIPAA Notice ({{Client Name}}))")
    generate = st.button("Generate Documents")

   if generate and excel_file and output_name_format:
    df = pd.read_excel(excel_file)

    if df.empty:
        st.error("⚠️ Your Excel file has no rows. Please check the file and try again.")
        st.stop()

    # Show preview
    st.subheader("🔍 Preview First Row of Excel Data")
    st.dataframe(df.head(1))

    # Show columns for debugging
    st.markdown("**Columns in Excel:**")
    st.code(", ".join(df.columns))

    # Preview filename
    preview_filename = output_name_format
    for key, val in df.iloc[0].items():
        preview_filename = preview_filename.replace(f"{{{{{key}}}}}", str(val))
    st.markdown("**📄 Preview Filename for First Row:**")
    st.code(preview_filename)

    # Generate documents
    left, right = "{{", "}}"
    with tempfile.TemporaryDirectory() as temp_dir:
        word_dir = os.path.join(temp_dir, "Word Documents")
        os.makedirs(word_dir)

        for idx, row in df.iterrows():
            doc = Document(template_path)

            for para in doc.paragraphs:
                for key, val in row.items():
                    if pd.api.types.is_datetime64_any_dtype([val]) or isinstance(val, pd.Timestamp):
                        val = val.strftime("%-m/%-d/%Y")
                    placeholder = f"{left}{key}{right}"
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(val))

            for table in doc.tables:
                for cell in table._cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for key, val in row.items():
                                if pd.api.types.is_datetime64_any_dtype([val]) or isinstance(val, pd.Timestamp):
                                    val = val.strftime("%-m/%-d/%Y")
                                placeholder = f"{left}{key}{right}"
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(val))

            name_for_file = output_name_format
            for key, val in row.items():
                if pd.api.types.is_datetime64_any_dtype([val]) or isinstance(val, pd.Timestamp):
                    val = val.strftime("%-m/%-d/%Y")
                name_for_file = name_for_file.replace(f"{left}{key}{right}", str(val))
            filename = name_for_file + ".docx"

            doc_path = os.path.join(word_dir, filename)
            doc.save(doc_path)

        # Zip files
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zip_out:
            for file in os.listdir(word_dir):
                full_path = os.path.join(word_dir, file)
                arcname = os.path.join("Word Documents", file)
                zip_out.write(full_path, arcname=arcname)

        st.success("✅ Word documents generated!")
        st.download_button(
            label="📦 Download All (Word Only – PDF not supported on Streamlit Cloud)",
            data=zip_buffer.getvalue(),
            file_name="word_documents.zip",
            mime="application/zip"
        )

elif tool == "📖 Instructions & Support":
    st.header("📘 Instructions")
    st.markdown("""
    Fill in the applicable fields:

    ### 📂 Demands:
    - Client Name
    - Incident Date
    - Summary
    - Damages

    ### 📨 FOIA:
    - Client ID
    - Case Synopsis
    - Potential Requests
    - Case Type
    - Facility or System
    - Defendant Role

    Click **Generate** to create your letters and download the results.
    """)

    st.subheader("🐞 Report a Bug")
    with st.form("report_form"):
        issue = st.text_area("Describe the issue:")
        submitted = st.form_submit_button("Submit")
        if submitted:
            with open("error_reports.txt", "a", encoding="utf-8") as f:
                f.write(issue + "\n---\n")
            st.success("✅ Issue submitted. Thank you!")

else:
    st.warning("🚧 This section is currently under development.")

st.markdown("""
<hr style="margin-top: 2rem;">
<div style="text-align: center; font-size: 0.85rem; color: gray;">
&copy; 2025 Stinar Gould Grieco & Hensley. All rights reserved.
</div>
""", unsafe_allow_html=True)

